import json
from docx import Document
from pandas import read_excel

base_path = "base/"
file_path = "result/"
info_list = []

# my_file = Path(file_path)
# with open("info_list.json","r", encoding="utf-8") as f:
#     info_list = json.loads(f.read())

# for info in info_list:
#     document = Document(base_path + 'work_doc.docx')
#     for p in document.paragraphs:
#         for keys in info.keys():
#             if "{{{%s}}}" % keys in p.text:
#                 p.text = p.text.replace("{{{%s}}}" % keys, info[keys])
#         document.save(file_path + 'complete_%s.docx' % (info.get("num")))

my_sheet = 'Sheet1'
file_name = 'info_list.xlsx' # name of your excel file
info_list = json.loads(read_excel(file_name, sheet_name = my_sheet).to_json(orient='records'))
for info in info_list:
    document = Document(base_path + 'work_doc.docx')
    for p in document.paragraphs:
        for keys in info.keys():
            if "{{{%s}}}" % keys in p.text:
                p.text = p.text.replace("{{{%s}}}" % keys, str(info[keys]))
        document.save(file_path + 'complete_%s.docx' % (info.get("num")))
